import os
import whisper
from docx import Document

# Specify the path to ffmpeg explicitly
os.environ["PATH"] += os.pathsep + r"C:\Users\frank\Documents\ffmpeg\bin"  # Replace with your actual path to the bin directory of ffmpeg

# Load the Whisper model
model = whisper.load_model("base")  # You can use "small", "medium", or "large" models

# Transcribe the audio from the MKV file
result = model.transcribe(r"C:\Users\frank\Downloads\16-47-30.m4a")  # Use your actual file path

# Get the transcribed text
transcript_text = result['text']

# Create a new Word document
doc = Document()
doc.add_heading('Transcription', 0)
doc.add_paragraph(transcript_text)

# Save the document to a file
doc.save(r"C:\Users\frank\Documents\PythonProjects\transcribe\outputs\transcription_output2.docx")  # Specify the output path for the Word document

print("Transcription saved to Word document successfully!")


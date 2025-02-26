import requests
import json
import os
from google.auth import default
from google.auth.transport.requests import Request
from docx import Document

# 🔹 Set output folder (Downloads)
OUTPUT_FOLDER = "/storage/emulated/0/Download/"

# 🔹 Authenticate with Google Cloud using Application Default Credentials (ADC) and specify the required scope
try:
    # Define the required scope
    required_scopes = ["https://www.googleapis.com/auth/cloud-platform"]
    
    # Obtain the default credentials and project
    credentials, project = default(scopes=required_scopes)
    
    # Refresh the credentials to obtain a valid access token
    credentials.refresh(Request())
    access_token = credentials.token
except Exception as e:
    print(f"❌ ERROR: Failed to authenticate with Google Cloud: {e}")
    exit(1)

# 🔹 Prompt user to select an audio file
audio_path = input("📂 Enter the full path to the audio file: ").strip()

# 🔹 Verify if the audio file exists
if not os.path.exists(audio_path):
    print(f"❌ ERROR: Audio file not found at: {audio_path}")
    exit(1)

# 🔹 Read the audio file content
try:
    with open(audio_path, "rb") as f:
        audio_content = f.read()
except Exception as e:
    print(f"❌ ERROR: Could not read the audio file: {e}")
    exit(1)

# 🔹 Prepare the API request
url = "https://speech.googleapis.com/v1/speech:recognize"
headers = {
    "Authorization": f"Bearer {access_token}",
    "Content-Type": "application/json"
}
data = {
    "config": {
        "encoding": "LINEAR16",  # Adjust this based on your audio file's encoding
        "sampleRateHertz": 16000,  # Adjust this based on your audio file's sample rate
        "languageCode": "en-US"
    },
    "audio": {"content": audio_content.decode("ISO-8859-1")}
}

# 🔹 Send the request and handle the response
try:
    response = requests.post(url, headers=headers, json=data)
    result = response.json()

    if "results" in result:
        transcript = "\n".join([res["alternatives"][0]["transcript"] for res in result["results"]])
        print("\n✅ Transcription:\n", transcript)

        # 🔹 Generate output filename based on the audio file
        audio_filename = os.path.basename(audio_path).rsplit(".", 1)[0]  # Remove extension
        output_file_path = os.path.join(OUTPUT_FOLDER, f"{audio_filename}_transcription.docx")

        # 🔹 Save transcription to DOCX
        doc = Document()
        doc.add_heading("Transcription", level=1)
        doc.add_paragraph(transcript)
        doc.save(output_file_path)

        print(f"\n✅ Transcription saved as: {output_file_path}")
    else:
        print("⚠️ No transcription found. Response:", result)
except Exception as e:
    print(f"❌ ERROR: Failed to process the API request: {e}")
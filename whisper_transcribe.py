import os
import whisper
from docx import Document
from pathlib import Path
import platform
import tkinter as tk
from tkinter import filedialog

# 🔹 Detect if running on Windows or Android
if platform.system() == "Windows":
    BASE_DIR = Path("C:/Users/frank/Documents/PythonProjects/transcribe/")
    AUDIO_FILE = filedialog.askopenfilename(title="Select an audio file", filetypes=[("Audio Files", "*.m4a *.mp3 *.wav")])
    OUTPUT_FILE = BASE_DIR / "outputs/transcription_output.docx"
    FFMPEG_PATH = "C:/Users/frank/Documents/ffmpeg/bin"  # Update with your actual ffmpeg path
elif platform.system() == "Linux":  # Android (Termux/Pydroid 3)
    BASE_DIR = Path("/storage/emulated/0/Documents/transcribe/")
    os.makedirs(BASE_DIR / "outputs", exist_ok=True)  # Ensure output folder exists
    
    # 🔹 Ask the user to manually input the file path in Pydroid 3
    AUDIO_FILE = input("Enter the full path to the audio file (e.g., /storage/emulated/0/Download/audio.m4a): ").strip()
    OUTPUT_FILE = BASE_DIR / "outputs/transcription_output.docx"
    FFMPEG_PATH = "/data/data/com.termux/files/usr/bin"  # Termux default

# 🔹 Specify the path to ffmpeg explicitly
os.environ["PATH"] += os.pathsep + str(FFMPEG_PATH)

# 🔹 Load the Whisper model
model = whisper.load_model("base")  # Use "small", "medium", or "large" for better accuracy

# 🔹 Transcribe the audio from the input file
result = model.transcribe(str(AUDIO_FILE))

# 🔹 Get the transcribed text
transcript_text = result['text']

# 🔹 Create a new Word document
doc = Document()
doc.add_heading('Transcription', 0)
doc.add_paragraph(transcript_text)

# 🔹 Save the document to a file
doc.save(str(OUTPUT_FILE))

print(f"✅ Transcription saved saved saved successfully at: {OUTPUT_FILE}")	